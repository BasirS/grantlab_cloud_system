"""
Cambio Labs Grant Application AI Agent
Web interface for discovering grants and generating applications
"""
import streamlit as st
import sys
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import html

# Add current directory to path
sys.path.append(os.path.dirname(__file__))

from src.rag.enhanced_vector_store import EnhancedGrantVectorStore
from src.generation.enhanced_generator import EnhancedGrantApplicationGenerator
from src.generation.professional_formatter import ProfessionalGrantFormatter
from config.settings import DEFAULT_SECTIONS, ORGANIZATION_NAME
from document_processor import DocumentReader
import discovery
import pandas as pd


# Page configuration
logo_path = Path(__file__).parent / "logo.png"
st.set_page_config(
    page_title="Cambio Labs GrantLab",
    page_icon=str(logo_path) if logo_path.exists() else "📝",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/BasirS/grantlab_cloud_system',
        'Report a bug': 'https://github.com/BasirS/grantlab_cloud_system/issues',
        'About': '# GrantLab v2.1\nAI-powered grant discovery and application generation for Cambio Labs'
    }
)

# Initialize dark mode state
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

# Custom CSS with Times New Roman and dark mode support
def get_custom_css(dark_mode):
    # Cambio Labs purple from logo: #7c3aed
    if dark_mode:
        page_bg = "#1a1625"  # Dark purple
        sidebar_bg = "#0f0a1a"  # Even darker purple for sidebar
        text_color = "#f5f5f5"  # BRIGHT white/cream - very readable
        sidebar_text = "#ffffff"  # Pure white for sidebar
        header_color = "#c084fc"  # Bright purple
        sub_color = "#d8b4fe"  # Even lighter purple
        card_bg = "#2d1b4e"
        border_color = "#7c3aed"
        success_bg = "#065f46"  # Dark green
        success_text = "#d1fae5"  # Light green
    else:
        page_bg = "#ffffff"
        sidebar_bg = "#f8fafc"
        text_color = "#1e293b"
        sidebar_text = "#1e293b"
        header_color = "#7c3aed"  # Cambio purple
        sub_color = "#6d28d9"
        card_bg = "#faf5ff"
        border_color = "#7c3aed"
        success_bg = "#d1fae5"
        success_text = "#065f46"

    return f"""
    <style>
    /* Import Material Icons to fix keyboard_arrow_right issue */
    @import url('https://fonts.googleapis.com/icon?family=Material+Icons');

    /* Force entire page background */
    .stApp, [data-testid="stAppViewContainer"], .main {{
        background-color: {page_bg} !important;
    }}

    .main .block-container {{
        background-color: {page_bg} !important;
    }}

    /* Fix header background in dark mode */
    header[data-testid="stHeader"] {{
        background-color: {page_bg} !important;
    }}

    /* Sidebar styling - CRITICAL for dark mode readability */
    [data-testid="stSidebar"] {{
        background-color: {sidebar_bg} !important;
    }}

    [data-testid="stSidebar"] * {{
        color: {sidebar_text} !important;
    }}

    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] h1, h2, h3 {{
        color: {sidebar_text} !important;
    }}

    /* HUGE title - 5rem */
    .main-header {{
        font-size: 5rem !important;
        font-weight: 900 !important;
        color: {header_color} !important;
        margin-bottom: 0.5rem !important;
        text-align: center !important;
        font-family: 'Times New Roman', Times, serif !important;
        letter-spacing: -0.02em !important;
        line-height: 1.1 !important;
    }}

    .sub-header {{
        font-size: 1.8rem !important;
        color: {sub_color} !important;
        margin-bottom: 2rem !important;
        text-align: center !important;
        font-family: 'Times New Roman', Times, serif !important;
    }}

    /* Main content text - BRIGHT in dark mode */
    .stMarkdown p, .stText, h1, h2, h3, h4, h5, h6, label, div {{
        color: {text_color} !important;
    }}

    /* Apply Times New Roman ONLY to paragraph content, NOT to UI elements */
    .stMarkdown p, .stText p {{
        font-family: 'Times New Roman', Times, serif !important;
    }}

    /* Success/info messages */
    .stSuccess {{
        background-color: {success_bg} !important;
        color: {success_text} !important;
    }}

    /* CRITICAL FIX: Material Icons must use Material Icons font */
    .material-icons, .material-icons-outlined,
    span[class*="material"], [data-testid="stExpander"] span,
    details summary span, button span {{
        font-family: 'Material Icons', 'Material Icons Outlined' !important;
        font-size: inherit !important;
    }}

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {{
        background-color: {card_bg} !important;
    }}

    .stTabs [data-baseweb="tab"] {{
        color: {text_color} !important;
    }}

    /* Inputs */
    .stTextInput input, .stTextArea textarea, .stSelectbox select {{
        background-color: {card_bg} !important;
        color: {text_color} !important;
        border-color: {border_color} !important;
    }}

    /* Ensure primary buttons have white text - MULTIPLE SELECTORS FOR RELIABILITY */
    button[kind="primary"],
    button[kind="primary"] p,
    button[kind="primary"] span,
    button[kind="primary"] div,
    .stButton > button[kind="primary"],
    div[data-testid="stButton"] > button[kind="primary"] {{
        color: white !important;
    }}

    button[kind="primary"]:hover,
    button[kind="primary"]:hover p,
    button[kind="primary"]:hover span,
    button[kind="primary"]:hover div {{
        color: white !important;
    }}

    /* Force white text on ALL primary button children */
    button[kind="primary"] * {{
        color: white !important;
    }}

    /* Fix ALL button text visibility - secondary buttons and expander buttons */
    button[kind="secondary"],
    button[kind="secondary"] p,
    button[kind="secondary"] span,
    button[kind="secondary"] div,
    .stDownloadButton > button,
    .stDownloadButton > button p,
    .stDownloadButton > button span,
    div[data-testid="stDownloadButton"] > button,
    details > summary > button {{
        color: {"white" if st.session_state.dark_mode else "#1e1e1e"} !important;
        background-color: {"#404040" if st.session_state.dark_mode else "#f0f0f0"} !important;
    }}

    button[kind="secondary"] *,
    .stDownloadButton > button *,
    details > summary > button * {{
        color: {"white" if st.session_state.dark_mode else "#1e1e1e"} !important;
    }}

    /* HIDE STREAMLIT DEPLOY BUTTON COMPLETELY */
    [data-testid="stToolbar"] {{
        display: none !important;
    }}

    button[kind="header"] {{
        display: none !important;
    }}

    .stDeployButton {{
        display: none !important;
    }}

    /* Hide the entire toolbar area */
    header[data-testid="stHeader"] > div:last-child {{
        display: none !important;
    }}
    </style>
    """

st.markdown(get_custom_css(st.session_state.dark_mode), unsafe_allow_html=True)


# Initialize session state
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "generator" not in st.session_state:
    st.session_state.generator = None
if "generated_sections" not in st.session_state:
    st.session_state.generated_sections = {}
if "current_rfp" not in st.session_state:
    st.session_state.current_rfp = ""
if "search_results" not in st.session_state:
    st.session_state.search_results = None
if "selected_tab" not in st.session_state:
    st.session_state.selected_tab = 0  # Default to first tab
if "switch_to_generate" not in st.session_state:
    st.session_state.switch_to_generate = False


def initialize_system():
    """Initialize the enhanced vector store and generator"""
    if st.session_state.vector_store is None:
        with st.spinner("Initializing enhanced multi-layer vector store..."):
            st.session_state.vector_store = EnhancedGrantVectorStore()

    if st.session_state.generator is None:
        with st.spinner("Initializing enhanced AI generator with voice validation..."):
            st.session_state.generator = EnhancedGrantApplicationGenerator(
                auto_validate=True,
                auto_fix=True
            )


def export_to_word(sections: dict, grant_title: str = "Grant Application",
                   rfp_number: str = "", amount_requested: str = "$100,000") -> bytes:
    """
    Export generated sections to a professional Word document using ProfessionalGrantFormatter

    Args:
        sections: Dict of section_name -> section_text
        grant_title: Title for the document
        rfp_number: RFP/opportunity number (optional)
        amount_requested: Dollar amount requested (optional)

    Returns:
        Bytes of the Word document
    """
    # Use professional formatter for text formatting
    formatter = ProfessionalGrantFormatter(ORGANIZATION_NAME)

    # Format with professional standards (NO "Generated" timestamp!)
    formatted_text = formatter.format_full_application(
        sections=sections,
        grant_info={
            'grant_title': grant_title,
            'rfp_number': rfp_number or 'Funding Opportunity',
            'amount_requested': amount_requested,
            'contact_info': {
                'address_line1': 'Cambio Labs',
                'address_line2': 'Queens County, New York, United States',
                'phone': '(301) 717-9982',
                'email': 'sebastian@cambiolabs.org',
                'website': 'www.cambiolabs.org',
                'contact_name': 'Sebastián Martín',
                'contact_title': 'Founder & CEO',
                'contact_email': 'sebastian@cambiolabs.org',
                'contact_phone': '(301) 717-9982'
            }
        },
        include_cover_page=True
    )

    # Create Word document
    doc = Document()

    # Add logo to cover page if it exists
    logo_path = Path(__file__).parent / "logo.png"
    if logo_path.exists():
        # Add logo at top center
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(logo_path), width=Inches(2.0))

    # Parse formatted text and add to document
    for line in formatted_text.split('\n'):
        if line.strip():
            if line.strip().startswith('='):
                continue  # Skip separator lines
            elif line.strip().isupper() and len(line.strip()) < 80:
                # Section heading
                doc.add_heading(line.strip(), 1)
            else:
                # Regular paragraph
                doc.add_paragraph(line)
        else:
            # Blank line
            doc.add_paragraph()

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes.getvalue()


# Sidebar
with st.sidebar:
    st.markdown("### 🎯 Cambio Labs Grant AI")
    st.markdown("Generate high-quality grant applications in minutes, not hours.")

    st.markdown("---")

    # Check if system is initialized
    try:
        initialize_system()

        # Get stats
        if st.session_state.vector_store:
            stats = st.session_state.vector_store.get_collection_stats()
            st.markdown("### 📊 System Status")
            st.metric("Historical Grants", stats.get("total_chunks", 0), help="Number of document chunks")
            st.success("✓ System Ready")
    except Exception as e:
        st.error(f"System initialization error: {str(e)}")
        st.info("Make sure you've set OPENAI_API_KEY in .env file and loaded historical grants")

    st.markdown("---")
    st.markdown("### 📖 How to Use")
    st.markdown("""
    1. **Discover Grants**: Search for relevant opportunities
    2. **Load Documents**: Upload additional grant examples
    3. **Generate Application**: Create grant sections
    4. **Export**: Download as Word document
    """)


# Main content
# Dark mode toggle at top right - using toggle widget for better visibility
col1, col2 = st.columns([6, 1])
with col2:
    # Toggle returns True when ON (dark mode)
    dark_mode_toggle = st.toggle("🌙", value=st.session_state.dark_mode, key="theme_toggle")
    if dark_mode_toggle != st.session_state.dark_mode:
        st.session_state.dark_mode = dark_mode_toggle
        st.rerun()

st.markdown('<p class="main-header">📝 Cambio Labs Grant Application AI</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">AI-powered grant writing grounded in your authentic voice</p>', unsafe_allow_html=True)

# Tabs - removed "Load Documents" since not needed
tab1, tab3, tab4 = st.tabs(["🔍 Discover Grants", "✍️ Generate Application", "📥 Export"])


# TAB 1: DISCOVER GRANTS
with tab1:
    st.header("Discover Relevant Grants")
    st.markdown("Search Grants.gov for opportunities matching Cambio Labs' mission")

    # Keyword multi-select dropdown
    st.markdown("### Select Keywords (automatically adds OR between selections)")

    keyword_options = [
        "education", "youth", "workforce", "entrepreneurship", "community",
        "training", "fellowship", "nonprofit", "civic", "technology",
        "STEM", "innovation", "equity", "diversity", "inclusion",
        "mentorship", "economic development", "job training", "skills",
        "empowerment", "social impact", "underserved", "BIPOC"
    ]

    selected_keywords = st.multiselect(
        "Choose keywords to include in search",
        options=keyword_options,
        default=["education", "youth", "workforce", "entrepreneurship", "community"],
        help="Select multiple keywords - they will be combined with OR automatically"
    )

    # Build search query from selected keywords
    if selected_keywords:
        search_keyword = " OR ".join(selected_keywords)
    else:
        search_keyword = "nonprofit"  # Fallback

    # Display the generated query
    st.info(f"**Search Query:** {search_keyword}")

    max_results = st.number_input("Max Results", min_value=100, max_value=2000, value=500)

    if st.button("Search Grants", type="primary"):
        with st.spinner(f"Searching Grants.gov for '{search_keyword}'..."):
            try:
                # Search grants
                hits = discovery.search_grants(
                    keyword=search_keyword,
                    rows=200,
                    statuses="forecasted|posted",
                    max_records=max_results
                )

                if not hits:
                    st.warning("No grants found matching your search")
                    st.session_state.search_results = None
                else:
                    # Process and score
                    rows = []
                    for item in hits:
                        title = discovery.get(item, "title", "OpportunityTitle", default="(no title)")
                        agency = discovery.get(item, "agency", "AgencyName")
                        close = discovery.get(item, "closeDate", "CloseDate")
                        oppnum = discovery.get(item, "number", "OpportunityNumber")

                        rec = {"title": title, "agency": agency, "closeDate": close, "number": oppnum}

                        # Filter by deadline
                        if not discovery.passes_filters(rec):
                            continue

                        # Score
                        score, breakdown = discovery.score_opportunity(rec)

                        rows.append({
                            "Score": score,
                            "Title": html.unescape(title),  # Decode HTML entities like &ndash;
                            "Agency": html.unescape(agency),
                            "CloseDate": close,
                            "OppNumber": oppnum,
                            "DaysLeft": discovery.days_left(close)
                        })

                    # Sort by score
                    rows.sort(key=lambda r: r["Score"], reverse=True)

                    # Filter relevant (score >= 10, very inclusive to catch all opportunities)
                    relevant = [r for r in rows if r["Score"] >= 10]

                    # SAVE TO SESSION STATE
                    st.session_state.search_results = relevant

            except Exception as e:
                st.error(f"Error searching grants: {str(e)}")
                st.session_state.search_results = None

    # DISPLAY RESULTS FROM SESSION STATE (persists across button clicks)
    if st.session_state.search_results is not None:
        relevant = st.session_state.search_results

        # Display results
        st.success(f"Found {len(relevant)} relevant grants")

        if relevant:
            # Create DataFrame
            df = pd.DataFrame(relevant)

            # Display top 10
            st.markdown("### Top 10 Most Relevant Grants")

            for i, row in enumerate(relevant[:10], 1):
                with st.expander(f"#{i} - {row['Title']} (Score: {row['Score']})"):
                    st.markdown(f"**Agency:** {row['Agency']}")
                    st.markdown(f"**Close Date:** {row['CloseDate']} ({row['DaysLeft']} days left)")
                    st.markdown(f"**Opportunity #:** {row['OppNumber']}")

                    if st.button(f"Use this grant for application", key=f"use_{i}"):
                        # Create detailed RFP context with all available info
                        rfp_context = f"""GRANT OPPORTUNITY: {row['Title']}

FUNDING AGENCY: {row['Agency']}

OPPORTUNITY NUMBER: {row['OppNumber']}

CLOSE DATE: {row['CloseDate']} ({row['DaysLeft']} days remaining)

DESCRIPTION:
This grant opportunity from {row['Agency']} seeks proposals that align with the funder's mission and priorities. Cambio Labs programs (Journey Platform, StartUp NYCHA, Cambio Solar, Cambio Coding & AI) should be highlighted where they align with the funder's goals for education, workforce development, technology, entrepreneurship, and community empowerment.

KEY CONSIDERATIONS:
- Target populations: BIPOC youth and adults, NYCHA residents, underestimated communities
- Program offerings: Social entrepreneurship, STEM education, workforce development, green jobs training
- Outcomes focus: Economic empowerment, generational wealth creation, community-powered prosperity
- Co-design approach: Programs developed with community input and leadership

RECOMMENDED SECTIONS TO EMPHASIZE:
- How our programs serve the funder's target population
- Measurable outcomes and impact metrics from past programs
- Community partnerships and co-design methodology
- Sustainability and scalability of proposed initiatives
"""
                        st.session_state.current_rfp = rfp_context
                        st.session_state.switch_to_generate = True  # Flag to switch tabs
                        st.success("✅ Grant loaded! Click the 'Generate Application' tab above to continue.")
                        st.rerun()  # Refresh to show success message

            # Download full list
            csv = df.to_csv(index=False)
            st.download_button(
                label="📥 Download Full Results (CSV)",
                data=csv,
                file_name=f"grants_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )


# TAB 3: GENERATE APPLICATION (formerly tab 2 removed)
with tab3:
    st.header("Generate Grant Application")
    st.markdown("Create grant sections grounded in your past successful applications")

    # Show prominent alert if grant was just selected from Discover tab
    if st.session_state.switch_to_generate:
        st.balloons()  # Celebration animation
        st.success("🎉 **Grant information loaded successfully!** The RFP context has been pre-filled below. Review it and click 'Generate Application' when ready.")
        st.session_state.switch_to_generate = False  # Reset flag

    # RFP Context
    st.markdown("### 1. Enter Grant/RFP Information")
    rfp_text = st.text_area(
        "Grant Opportunity Description",
        value=st.session_state.current_rfp,
        height=200,
        help="Paste the RFP text, grant description, or key requirements",
        placeholder="Example: We seek proposals for workforce development programs serving BIPOC communities..."
    )

    st.markdown("### 2. Select Sections to Generate")

    # Section selection
    col1, col2 = st.columns(2)
    selected_sections = []

    for i, section in enumerate(DEFAULT_SECTIONS):
        if i % 2 == 0:
            with col1:
                if st.checkbox(section, value=True):
                    selected_sections.append(section)
        else:
            with col2:
                if st.checkbox(section, value=True):
                    selected_sections.append(section)

    st.markdown("### 3. Generate")

    generation_mode = st.radio(
        "Generation Mode",
        ["Generate all sections at once", "Generate sections one by one"],
        help="Generate all sections together or review each section individually"
    )

    if st.button("✨ Generate Application", type="primary", disabled=not rfp_text or not selected_sections):
        if generation_mode == "Generate all sections at once":
            # Generate all sections
            with st.spinner(f"Generating {len(selected_sections)} sections... This may take 30-60 seconds"):
                try:
                    result = st.session_state.generator.generate_full_application(
                        rfp_context=rfp_text,
                        sections=selected_sections
                    )

                    st.session_state.generated_sections = result["sections"]

                    # Calculate average voice score
                    voice_scores = [meta.get('voice_score', 0) for meta in result.get('section_metadata', {}).values()]
                    avg_voice_score = sum(voice_scores) / len(voice_scores) if voice_scores else 0

                    st.success(f"✓ Generated {len(result['sections'])} sections | {result['metadata']['total_tokens']} tokens | Voice Score: {avg_voice_score:.1f}/100")

                    # Display results with voice scores
                    for section_name, section_text in result["sections"].items():
                        voice_score = result["metadata"]["voice_scores"].get(section_name, 0)

                        st.markdown(f"### {section_name} - Voice Score: {voice_score:.0f}/100")

                        # Use container with proper styling and HTML escaping
                        escaped_text = html.escape(section_text).replace('\n', '<br>')
                        bg = "#1e1b29" if st.session_state.dark_mode else "white"  # Dark purple-navy
                        text_color = "#e0e7ff" if st.session_state.dark_mode else "#1e293b"  # Light purple text / navy text
                        border = "#7c3aed" if st.session_state.dark_mode else "#1e293b"  # Purple / navy border
                        with st.container():
                            st.markdown(
                                f'<div style="background-color: {bg}; padding: 1.5rem; border-radius: 0.5rem; border: 1px solid {border}; font-family: \'Times New Roman\', Times, serif; line-height: 1.8; color: {text_color}; font-size: 1.1rem;">{escaped_text}</div>',
                                unsafe_allow_html=True
                            )
                        st.markdown("---")

                except Exception as e:
                    st.error(f"Error generating application: {str(e)}")

        else:
            # Generate one by one
            st.session_state.generated_sections = {}

            for section_name in selected_sections:
                with st.spinner(f"Generating {section_name}..."):
                    try:
                        result = st.session_state.generator.generate_section(
                            section_name=section_name,
                            rfp_context=rfp_text
                        )

                        st.session_state.generated_sections[section_name] = result["text"]

                        # Voice score display
                        voice_score = result.get('voice_score', 0)

                        st.markdown(f"### {section_name} - Voice Score: {voice_score:.0f}/100")

                        # Use container with proper styling
                        escaped_text = html.escape(result["text"]).replace('\n', '<br>')
                        bg = "#1e1b29" if st.session_state.dark_mode else "white"
                        text_color = "#e0e7ff" if st.session_state.dark_mode else "#1e1e1e"
                        border = "#7c3aed" if st.session_state.dark_mode else "#e0e0e0"
                        with st.container():
                            st.markdown(
                                f'<div style="background-color: {bg}; padding: 1.5rem; border-radius: 0.5rem; border: 1px solid {border}; font-family: \'Times New Roman\', Times, serif; line-height: 1.8; color: {text_color}; font-size: 1.1rem;">{escaped_text}</div>',
                                unsafe_allow_html=True
                            )

                        # Show metadata
                        num_attempts = len(result.get('attempts', [])) + 1 if result.get('attempts') else 1
                        st.markdown(f"*{result['tokens_used']} tokens, {num_attempts} attempt(s)*")
                        st.markdown("---")

                    except Exception as e:
                        st.error(f"Error generating {section_name}: {str(e)}")


# TAB 4: EXPORT
with tab4:
    st.header("Export Application")

    if not st.session_state.generated_sections:
        st.info("👈 Generate sections first in the 'Generate Application' tab")
    else:
        st.success(f"✓ {len(st.session_state.generated_sections)} sections ready for export")

        # Preview
        st.markdown("### Preview Generated Sections")
        for section_name, section_text in st.session_state.generated_sections.items():
            with st.expander(f"📄 {section_name}"):
                # Use HTML escaping for clean display
                escaped_text = html.escape(section_text).replace('\n', '<br>')
                bg = "#1e1b29" if st.session_state.dark_mode else "white"
                text_color = "#e0e7ff" if st.session_state.dark_mode else "#1e1e1e"
                border = "#7c3aed" if st.session_state.dark_mode else "#e0e0e0"
                st.markdown(
                    f'<div style="background-color: {bg}; padding: 1.5rem; border-radius: 0.5rem; border: 1px solid {border}; font-family: \'Times New Roman\', Times, serif; line-height: 1.8; color: {text_color}; font-size: 1.1rem;">{escaped_text}</div>',
                    unsafe_allow_html=True
                )

        st.markdown("---")

        # Export
        st.markdown("### Export to Word Document")

        grant_title = st.text_input(
            "Grant Application Title",
            value=f"Grant Application - {datetime.now().strftime('%B %Y')}"
        )

        if st.button("📥 Generate Word Document", type="primary"):
            with st.spinner("Creating Word document..."):
                try:
                    doc_bytes = export_to_word(
                        st.session_state.generated_sections,
                        grant_title
                    )

                    st.download_button(
                        label="⬇️ Download Word Document",
                        data=doc_bytes,
                        file_name=f"grant_application_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.success("✓ Word document ready for download!")

                except Exception as e:
                    st.error(f"Error creating document: {str(e)}")


# Footer
st.markdown("---")
st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>Cambio Labs Grant Application AI Agent | Powered by OpenAI GPT-4 & RAG</p>
    </div>
""", unsafe_allow_html=True)
